"""
Bus Pass Management System - Documentation Generator
Generates PDF and DOCX documentation with proper formatting.

Requirements: pip install fpdf2 python-docx Pillow
Usage: python generate_documentation.py

Output:
  - Bus_Pass_Management_System_Documentation.pdf
  - Bus_Pass_Management_System_Documentation.docx
"""

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SS_DIR = os.path.join(BASE_DIR, "Bus-Pass-Management-System-SS")
OUTPUT_PDF = os.path.join(BASE_DIR, "Bus_Pass_Management_System_Documentation.pdf")
OUTPUT_DOCX = os.path.join(BASE_DIR, "Bus_Pass_Management_System_Documentation.docx")

# Screenshot mapping: (original filename, title, description, side: user/admin)
NARROW_NBSP = "\u202f"  # macOS uses narrow no-break space before AM/PM


def ss(name):
    """Fix screenshot filename - macOS uses narrow no-break space before AM/PM."""
    return name.replace(" AM.png", f"{NARROW_NBSP}AM.png")


SCREENSHOTS = [
    # User Side
    (ss("Screenshot 2026-03-22 at 11.22.51 AM.png"), "User Login Page",
     "The login page allows registered users to sign in using their email and password. Features include form validation, a 'Forgot Password' link, and a link to the registration page. URL: http://localhost:5173/login",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.22.32 AM.png"), "User Registration Page",
     "New users can create an account by providing their full name, email address, phone number, date of birth, and a strong password with confirmation. URL: http://localhost:5173/register",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.23.56 AM.png"), "Forgot Password Page",
     "Users can request a password reset code sent to their registered email. The page accepts an email address and sends a 6-character alphanumeric code valid for 15 minutes. URL: http://localhost:5173/forgot-password",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.25.19 AM.png"), "User Dashboard - Buy a Bus Pass",
     "The main dashboard displays source and destination dropdowns with 20 bus stops, real-time fare calculation for all pass types (Monthly, Quarterly, Half-Yearly, Yearly), and Buy Now buttons. URL: http://localhost:5173/dashboard",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.28.11 AM.png"), "User Dashboard - Active Pass with QR Code",
     "After admin approval, the dashboard shows the active bus pass with QR code for scanning, verification code for manual verification, validity dates, route details, and pass type information. URL: http://localhost:5173/dashboard",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.28.44 AM.png"), "Family Pass - Add Members (Step 1)",
     "Step 1 of the family pass purchase flow. Users can add multiple family members with the '+ Add Member' button. Shows a stepper navigation with 3 steps: Add Members, Route & Pass, Review. URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.29.11 AM.png"), "Family Pass - Add Family Member Dialog",
     "A modal dialog for adding a family member with fields for Full Name, Email Address, Date of Birth, and Relation (Spouse, Father, Mother, Son, Daughter, etc.). URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.29.27 AM.png"), "Family Pass - Member Added (Step 1)",
     "After adding a member, the family members list shows the added member with name, email, relation, and DOB. Members can be edited or removed. URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.29.54 AM.png"), "Family Pass - Route & Pass Selection (Step 2)",
     "Step 2 allows selecting source and destination stops with fare calculation per person. Shows all four pass types with prices and total calculation. URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.30.09 AM.png"), "Family Pass - Review & Confirm (Step 3)",
     "Final review step showing route details, pass type, family member list with individual pricing, and total amount with a 'Confirm & Buy' button. URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.31.29 AM.png"), "Family Pass - Active Passes with QR Codes",
     "After approval, family passes display with QR codes, verification codes, validity dates, member details, and route information for each family member. URL: http://localhost:5173/family-pass",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.31.55 AM.png"), "User Profile Page",
     "Displays the user's profile information including name, email, phone number, and date of birth in a clean read-only format. URL: http://localhost:5173/profile",
     "user"),
    (ss("Screenshot 2026-03-22 at 11.33.09 AM.png"), "User Dashboard - Mobile View",
     "Responsive mobile view of the user dashboard showing user details, active bus passes with QR codes, and pass information optimized for small screens. URL: http://localhost:5173/dashboard (mobile)",
     "user"),

    # Admin Side
    (ss("Screenshot 2026-03-22 at 11.26.20 AM.png"), "Admin Login Page",
     "Admin secure login portal with email and password fields. Displays 'Unauthorized access is prohibited' warning. URL: http://localhost:5174/login",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.26.55 AM.png"), "Admin Registration Page",
     "Admin registration form with fields for Full Name, Official Admin Email, Password, Confirm Password, and Access Level dropdown (Admin, Superadmin, Conductor). URL: http://localhost:5174/register",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.27.37 AM.png"), "Admin Dashboard - Pending Passes",
     "Admin dashboard showing welcome message, role badge, statistics cards (Total Passes, Active, Pending), and pending bus pass requests table with Approve/Reject action buttons. URL: http://localhost:5174/dashboard",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.33.49 AM.png"), "Admin Dashboard - Active Passes",
     "Active passes tab showing all approved passes with user details, member info, route, pass type, price, status, requested date, and validity dates. URL: http://localhost:5174/dashboard",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.34.03 AM.png"), "Admin Dashboard - Rejected Passes",
     "Rejected passes tab displaying all rejected pass applications with user information, route details, and rejection status. URL: http://localhost:5174/dashboard",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.37.42 AM.png"), "Admin Dashboard - QR Code Scanner & Manual Verification",
     "The verify bus pass section with a text input for 16-character code or JWT token, Verify button, and integrated camera-based QR code scanner for real-time pass scanning. URL: http://localhost:5174/dashboard",
     "admin"),
    (ss("Screenshot 2026-03-22 at 11.38.48 AM.png"), "Conductor Mode - Pass Verification Result",
     "Conductor verification view showing 'Pass Verified Successfully' with passenger name, phone, pass type, validity, route, family member details, fare paid, and active status. URL: http://localhost:5173/dashboard (conductor)",
     "admin"),
]


# ═══════════════════════════════════════════════════════
#  PDF Generation
# ═══════════════════════════════════════════════════════

class DocumentationPDF(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="mm", format="A4")
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        if self.page_no() > 2:
            self.set_font("Times", "I", 9)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8, "Bus Pass Management System - Project Documentation", align="C")
            self.ln(5)
            self.set_draw_color(0, 0, 0)
            self.set_line_width(0.3)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(5)

    def footer(self):
        if self.page_no() > 1:
            self.set_y(-20)
            self.set_font("Times", "I", 9)
            self.set_text_color(100, 100, 100)
            self.set_draw_color(0, 0, 0)
            self.set_line_width(0.3)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(3)
            self.cell(0, 10, f"Page {self.page_no() - 1}", align="C")

    def cover_page(self):
        self.add_page()
        self.ln(25)

        self.set_font("Times", "B", 18)
        self.set_text_color(0, 0, 0)
        self.cell(0, 10, "GLS UNIVERSITY", align="C", new_x="LMARGIN", new_y="NEXT")

        self.set_font("Times", "B", 14)
        self.cell(0, 8, "Faculty of Computer Applications & Information Technology", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(5)

        self.set_font("Times", "B", 16)
        self.cell(0, 10, "Integrated Master of Computer Applications [iMSc.IT]", align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Times", "B", 18)
        self.cell(0, 10, "Programme", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

        self.set_font("Times", "B", 16)
        self.cell(0, 10, "Semester VIII Project Report for", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

        self.set_font("Times", "B", 16)
        self.cell(0, 10, "221601801 FULL STACK DEVELOPMENT", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

        self.set_font("Times", "B", 16)
        self.cell(0, 12, '"Bus Pass Management System"', align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(10)

        self.set_font("Times", "", 14)
        self.cell(0, 10, "Submitted by:", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

        self.set_font("Times", "", 13)
        self.cell(0, 8, "Group No.:", align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, "Hardik Sanjaybhai Parmar - 202201619010054", align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, "Dhairy Pradipbhai Makadia - 202201619010152", align="C", new_x="LMARGIN", new_y="NEXT")

    def certificate_page(self):
        self.add_page()
        self.ln(15)

        self.set_font("Times", "B", 18)
        self.set_text_color(0, 0, 0)
        self.cell(0, 10, "GLS UNIVERSITY", align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Times", "", 14)
        self.cell(0, 8, "Faculty of Computer Applications & IT iMSc.IT Programme", align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, "Ahmedabad", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(10)

        self.set_font("Times", "B", 18)
        self.cell(0, 10, "CERTIFICATE", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(8)

        self.set_font("Times", "", 13)
        self.cell(0, 8, "This is to certify that", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

        self.set_font("Times", "B", 13)
        self.cell(0, 8, "1. Hardik Sanjaybhai Parmar", align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, "2. Dhairy Pradipbhai Makadia", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(5)

        self.set_font("Times", "", 12)
        cert_text = (
            "Students of Semester-VIII iMSc.IT, FCAIT, GLS University have successfully "
            "completed the Project of 221601801 FULL STACK DEVELOPMENT as a part of their "
            "curriculum for the academic year 2025-2026."
        )
        self.multi_cell(0, 7, cert_text, align="J")
        self.ln(30)

        self.set_font("Times", "", 12)
        self.cell(95, 8, "Date of Submission:", align="L")
        self.cell(95, 8, "Project Guide - Name & Sign", align="R", new_x="LMARGIN", new_y="NEXT")

    def add_heading(self, text, level=1):
        if level == 1:
            self.set_font("Times", "B", 16)
            self.set_text_color(0, 0, 0)
            self.ln(5)
            self.cell(0, 12, text, new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(0, 0, 0)
            self.set_line_width(0.5)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(6)
        elif level == 2:
            self.set_font("Times", "B", 14)
            self.set_text_color(0, 0, 0)
            self.ln(4)
            self.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
            self.ln(2)
        elif level == 3:
            self.set_font("Times", "B", 12)
            self.set_text_color(0, 0, 0)
            self.ln(2)
            self.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
            self.ln(2)

    def add_paragraph(self, text):
        self.set_font("Times", "", 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 6, text, align="J")
        self.ln(3)

    def add_bullet(self, text, indent=15):
        self.set_font("Times", "", 12)
        self.set_text_color(0, 0, 0)
        self.cell(indent)
        self.multi_cell(0, 6, f"-  {text}", align="L")
        self.ln(1)

    def add_table(self, headers, rows, col_widths=None):
        if col_widths is None:
            col_widths = [190 / len(headers)] * len(headers)

        # Header
        self.set_font("Times", "B", 11)
        self.set_fill_color(0, 0, 0)
        self.set_text_color(255, 255, 255)
        for i, h in enumerate(headers):
            self.cell(col_widths[i], 9, h, border=1, fill=True, align="C")
        self.ln()

        # Rows
        self.set_font("Times", "", 11)
        self.set_text_color(0, 0, 0)
        fill = False
        for row in rows:
            if fill:
                self.set_fill_color(235, 235, 235)
            else:
                self.set_fill_color(255, 255, 255)

            max_lines = 1
            for i, cell_text in enumerate(row):
                lines = self.multi_cell(col_widths[i], 7, str(cell_text), dry_run=True, output="LINES")
                max_lines = max(max_lines, len(lines))
            row_height = max_lines * 7

            y_start = self.get_y()
            x_start = self.get_x()
            if y_start + row_height > 270:
                self.add_page()
                y_start = self.get_y()

            for i, cell_text in enumerate(row):
                x_pos = x_start + sum(col_widths[:i])
                self.set_xy(x_pos, y_start)
                self.cell(col_widths[i], row_height, "", border=1, fill=fill)
                self.set_xy(x_pos + 1, y_start + 1)
                self.multi_cell(col_widths[i] - 2, 7, str(cell_text), align="L")

            self.set_xy(x_start, y_start + row_height)
            fill = not fill
        self.ln(5)

    def add_screenshot(self, filepath, title, description):
        self.ln(3)
        self.set_font("Times", "B", 12)
        self.set_text_color(0, 0, 0)
        self.cell(0, 8, f"Screenshot: {title}", new_x="LMARGIN", new_y="NEXT")

        if os.path.exists(filepath):
            if self.get_y() > 170:
                self.add_page()
            try:
                self.image(filepath, x=20, w=170)
            except Exception:
                self.set_font("Times", "I", 11)
                self.cell(0, 8, f"[Image could not be loaded]", new_x="LMARGIN", new_y="NEXT")
        else:
            y = self.get_y()
            self.set_draw_color(150, 150, 150)
            self.set_fill_color(245, 245, 245)
            self.rect(20, y, 170, 70, style="DF")
            self.set_xy(20, y + 28)
            self.set_font("Times", "I", 12)
            self.set_text_color(120, 120, 120)
            self.cell(170, 10, f"[Screenshot: {title}]", align="C")
            self.set_y(y + 75)

        self.set_font("Times", "I", 11)
        self.set_text_color(60, 60, 60)
        self.multi_cell(0, 6, f"Figure: {title} - {description}", align="J")
        self.ln(5)


def build_pdf():
    pdf = DocumentationPDF()

    # ── Cover Page ──
    pdf.cover_page()

    # ── Certificate Page ──
    pdf.certificate_page()

    # ── Table of Contents ──
    pdf.add_page()
    pdf.add_heading("TABLE OF CONTENTS")
    toc = [
        ("1.", "Introduction"),
        ("   I.", "Project Objective"),
        ("   II.", "Purpose"),
        ("   III.", "Modules"),
        ("2.", "System Architecture"),
        ("   I.", "Technology Stack"),
        ("   II.", "Architecture Diagram"),
        ("   III.", "Project Structure"),
        ("3.", "Features & Functionality"),
        ("   I.", "User Features"),
        ("   II.", "Admin Features"),
        ("   III.", "Authentication & Security"),
        ("4.", "API Endpoints"),
        ("5.", "Database Schema Design"),
        ("6.", "Business Logic"),
        ("   I.", "Fare Calculation"),
        ("   II.", "Bus Stops"),
        ("   III.", "Pass Lifecycle"),
        ("   IV.", "QR Code Verification"),
        ("7.", "Screenshots - User Side"),
        ("8.", "Screenshots - Admin Side"),
        ("9.", "Conclusion"),
    ]
    pdf.set_font("Times", "", 12)
    pdf.set_text_color(0, 0, 0)
    for num, title in toc:
        is_main = not num.startswith(" ")
        pdf.set_font("Times", "B" if is_main else "", 12)
        pdf.cell(15, 8, num)
        pdf.cell(165, 8, title, new_x="LMARGIN", new_y="NEXT")

    # ── 1. Introduction ──
    pdf.add_page()
    pdf.add_heading("1. INTRODUCTION")

    pdf.add_heading("I. Project Objective", level=2)
    pdf.add_paragraph(
        "The Bus Pass Management System is a comprehensive full-stack web application designed to "
        "digitize and streamline the process of issuing, managing, and verifying bus passes. The system "
        "replaces traditional paper-based bus pass systems with a modern, efficient, and secure digital "
        "solution that benefits both commuters and transport authorities."
    )
    pdf.add_paragraph(
        "The primary objective is to create a seamless experience for users to purchase bus passes online, "
        "enable administrators to manage pass approvals efficiently, and provide conductors with real-time "
        "QR-based verification capabilities. The system supports individual and family pass purchases with "
        "dynamic fare calculation based on distance between bus stops."
    )

    pdf.add_heading("II. Purpose", level=2)
    pdf.add_paragraph(
        "The purpose of this system is to modernize public transportation services by providing:"
    )
    pdf.add_bullet("Digital bus passes with QR code verification")
    pdf.add_bullet("Online application and approval system")
    pdf.add_bullet("Family pass management with auto-account creation")
    pdf.add_bullet("Real-time pass tracking and status updates")
    pdf.add_bullet("Role-based access for Admin, Superadmin, and Conductor")
    pdf.add_paragraph(
        "This system reduces waiting time at bus pass counters and ensures a smooth and secure "
        "experience for users and administrators."
    )

    pdf.add_heading("III. Modules", level=2)

    pdf.add_heading("1. User Authentication Module", level=3)
    pdf.add_bullet("User registration and login with form validation")
    pdf.add_bullet("Secure authentication using JWT (JSON Web Tokens)")
    pdf.add_bullet("Password encryption using Bcrypt")
    pdf.add_bullet("Forgot and reset password functionality via email")

    pdf.add_heading("2. Individual Pass Module", level=3)
    pdf.add_bullet("Select source and destination from 20 bus stops")
    pdf.add_bullet("Dynamic fare calculation using Haversine formula")
    pdf.add_bullet("Choose pass type (Monthly, Quarterly, Half-Yearly, Yearly)")
    pdf.add_bullet("Submit pass request for admin approval")

    pdf.add_heading("3. Family Pass Module", level=3)
    pdf.add_bullet("Add multiple family members (up to 10)")
    pdf.add_bullet("Auto-create user accounts for family members")
    pdf.add_bullet("Purchase passes for entire family in one transaction")
    pdf.add_bullet("3-step wizard: Add Members, Route & Pass, Review")

    pdf.add_heading("4. Dashboard Module", level=3)
    pdf.add_bullet("View all passes with status (Active, Pending, Expired, Rejected)")
    pdf.add_bullet("Display QR codes for active passes")
    pdf.add_bullet("Pass details with verification code")
    pdf.add_bullet("Responsive design for mobile and desktop")

    pdf.add_heading("5. Admin Module", level=3)
    pdf.add_bullet("Approve or reject pending pass applications")
    pdf.add_bullet("View passes filtered by status")
    pdf.add_bullet("Statistics dashboard (Total, Active, Pending counts)")
    pdf.add_bullet("Superadmin can manage other admin accounts")

    pdf.add_heading("6. Verification Module", level=3)
    pdf.add_bullet("QR code scanning using device camera")
    pdf.add_bullet("Manual verification using 16-character hex code")
    pdf.add_bullet("Real-time validation with pass details display")

    # ── 2. System Architecture ──
    pdf.add_page()
    pdf.add_heading("2. SYSTEM ARCHITECTURE")

    pdf.add_heading("I. Technology Stack", level=2)
    pdf.add_table(
        ["Component", "Technology"],
        [
            ["Frontend (User)", "React.js 19.x + Tailwind CSS + Vite"],
            ["Frontend (Admin)", "React.js 19.x + Tailwind CSS + Vite"],
            ["Backend", "Node.js + Express.js 5.x"],
            ["Database", "MongoDB Atlas (Mongoose 9.x ODM)"],
            ["Authentication", "JSON Web Token (JWT)"],
            ["Password Security", "Bcrypt.js"],
            ["QR Code", "qrcode library + html5-qrcode scanner"],
            ["Email Service", "Nodemailer (Gmail SMTP)"],
            ["Routing", "React Router 7.x"],
            ["Deployment", "Render.com (Cloud Platform)"],
        ],
        col_widths=[50, 140],
    )

    pdf.add_heading("II. Architecture Diagram", level=2)
    pdf.set_font("Courier", "B", 9)
    pdf.set_text_color(0, 0, 0)
    diagram = [
        "+---------------------+     +---------------------+",
        "|   User Frontend     |     |   Admin Frontend    |",
        "|   (React + Vite)    |     |   (React + Vite)    |",
        "|   Port: 5173        |     |   Port: 5174        |",
        "+----------+----------+     +----------+----------+",
        "           |                           |           ",
        "           +----------- + -------------+           ",
        "                        |                          ",
        "              +---------v---------+                ",
        "              |  Express.js API   |                ",
        "              |  (Node.js)        |                ",
        "              |  Port: 8000       |                ",
        "              |  JWT Auth + CORS  |                ",
        "              +---------+---------+                ",
        "                        |                          ",
        "              +---------v---------+                ",
        "              |   MongoDB Atlas   |                ",
        "              |   Collections:    |                ",
        "              |   users, admins,  |                ",
        "              |   buspasses       |                ",
        "              +-------------------+                ",
    ]
    for line in diagram:
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    pdf.add_heading("III. Project Structure", level=2)
    pdf.set_font("Courier", "", 9)
    structure = [
        "bus-pass-management-system/",
        "|-- backend/",
        "|   |-- controllers/     # Business logic handlers",
        "|   |-- models/          # Mongoose schemas",
        "|   |-- routes/          # API route definitions",
        "|   |-- middleware/      # Auth guards, HTTP logging",
        "|   |-- utils/           # Token, QR, email, logger",
        "|   |-- index.js         # Express server entry point",
        "|",
        "|-- frontend/            # Admin Portal (React + Vite)",
        "|   |-- src/auth/        # Login, Register, auth service",
        "|   |-- src/pages/       # Admin Dashboard",
        "|",
        "|-- user-frontend/       # User Portal (React + Vite)",
        "|   |-- src/auth/        # Login, Register, ForgotPassword",
        "|   |-- src/pages/       # Dashboard, FamilyPass, Profile",
        "|   |-- src/components/  # Navbar, Footer",
        "|   |-- src/layout/      # RootLayout wrapper",
    ]
    for line in structure:
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # ── 3. Features ──
    pdf.add_page()
    pdf.add_heading("3. FEATURES & FUNCTIONALITY")

    pdf.add_heading("I. User Features", level=2)
    pdf.add_bullet("User Registration with name, email, phone (10-digit), DOB, and strong password validation")
    pdf.add_bullet("JWT-based login with 1-day token expiry")
    pdf.add_bullet("Email-based password reset with 6-character code (15-min expiry)")
    pdf.add_bullet("Individual pass purchase with real-time fare display for 4 pass types")
    pdf.add_bullet("Family pass purchase for up to 10 members with auto-account creation")
    pdf.add_bullet("View pass history with QR codes, verification codes, and validity dates")
    pdf.add_bullet("Responsive design for desktop and mobile devices")

    pdf.add_heading("II. Admin Features", level=2)
    pdf.add_bullet("Pass approval/rejection workflow with QR code generation on approval")
    pdf.add_bullet("Camera-based QR scanner for real-time pass verification")
    pdf.add_bullet("Manual verification using 16-character hex code")
    pdf.add_bullet("Pass filtering by status (Pending, Active, Rejected, All)")
    pdf.add_bullet("Role-based access: Admin (manage passes), Conductor (verify only), Superadmin (full control)")
    pdf.add_bullet("Dashboard statistics: Total Passes, Active, Pending counts")

    pdf.add_heading("III. Authentication & Security", level=2)
    pdf.add_bullet("JWT token-based stateless authentication with 1-day expiration")
    pdf.add_bullet("Bcrypt password hashing with salt rounds")
    pdf.add_bullet("Role-based middleware: protectUser, protectAdmin, adminOnly, superAdminOnly, anyAdminRole")
    pdf.add_bullet("CORS whitelisting for development and production origins")
    pdf.add_bullet("Server-side input validation for all endpoints")

    # ── 4. API Endpoints ──
    pdf.add_page()
    pdf.add_heading("4. API ENDPOINTS")

    pdf.add_heading("User Authentication API", level=2)
    pdf.add_table(
        ["Method", "Endpoint", "Auth", "Description"],
        [
            ["POST", "/api/user/register", "Public", "Register new user"],
            ["POST", "/api/user/login", "Public", "User login, returns JWT"],
            ["POST", "/api/user/forgot-password", "Public", "Send reset code to email"],
            ["POST", "/api/user/reset-password", "Public", "Reset password with code"],
            ["GET", "/api/user/profile", "User", "Get user profile"],
            ["PUT", "/api/user/change-password", "User", "Change password"],
        ],
        col_widths=[20, 55, 18, 97],
    )

    pdf.add_heading("Bus Pass API (User)", level=2)
    pdf.add_table(
        ["Method", "Endpoint", "Auth", "Description"],
        [
            ["GET", "/api/user/bus-pass/stops", "Public", "Get 20 bus stops list"],
            ["GET", "/api/user/bus-pass/fare", "Public", "Calculate fare for route"],
            ["POST", "/api/user/bus-pass/buy", "User", "Purchase individual pass"],
            ["POST", "/api/user/bus-pass/buy-family", "User", "Purchase family passes"],
            ["GET", "/api/user/bus-pass/my-passes", "User", "Get user's passes"],
            ["POST", "/api/user/bus-pass/verify", "Public", "Verify pass (QR/code)"],
        ],
        col_widths=[20, 55, 18, 97],
    )

    pdf.add_heading("Admin API", level=2)
    pdf.add_table(
        ["Method", "Endpoint", "Auth", "Description"],
        [
            ["POST", "/api/admin/login", "Public", "Admin login"],
            ["POST", "/api/admin/register", "Public", "Register admin"],
            ["GET", "/api/admin/profile", "Admin", "Get admin profile"],
            ["GET", "/api/admin/bus-passes", "Admin", "Get passes by status"],
            ["PATCH", "/api/admin/bus-passes/:id/approve", "Admin", "Approve pass"],
            ["PATCH", "/api/admin/bus-passes/:id/reject", "Admin", "Reject pass"],
            ["POST", "/api/admin/bus-passes/verify", "Any", "Verify pass"],
        ],
        col_widths=[20, 60, 18, 92],
    )

    # ── 5. Database Schema ──
    pdf.add_page()
    pdf.add_heading("5. DATABASE SCHEMA DESIGN")

    pdf.add_heading("User Collection", level=2)
    pdf.add_table(
        ["Field", "Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key", "Auto-generated unique ID"],
            ["name", "String", "Required", "User's full name"],
            ["email", "String", "Required, Unique", "Email address"],
            ["phone", "String", "Required", "10-digit phone number"],
            ["dateOfBirth", "Date", "Required", "Date of birth"],
            ["password", "String", "Required", "Bcrypt hashed password"],
            ["isActive", "Boolean", "Default: true", "Account status"],
        ],
        col_widths=[30, 22, 45, 93],
    )

    pdf.add_heading("Admin Collection", level=2)
    pdf.add_table(
        ["Field", "Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key", "Auto-generated unique ID"],
            ["name", "String", "Required", "Admin's full name"],
            ["email", "String", "Required, Unique", "Email address"],
            ["password", "String", "Required", "Bcrypt hashed password"],
            ["role", "String", "Enum", "admin / superadmin / conductor"],
            ["isActive", "Boolean", "Default: true", "Account status"],
        ],
        col_widths=[30, 22, 45, 93],
    )

    pdf.add_heading("BusPass Collection", level=2)
    pdf.add_table(
        ["Field", "Type", "Constraints", "Description"],
        [
            ["_id", "ObjectId", "Primary Key", "Auto-generated unique ID"],
            ["user", "ObjectId", "Ref: User", "Pass owner reference"],
            ["passType", "String", "Enum", "monthly/quarterly/half-yearly/yearly"],
            ["price", "Number", "Required", "Calculated fare amount"],
            ["source", "String", "Required", "Source bus stop"],
            ["destination", "String", "Required", "Destination bus stop"],
            ["validFrom", "Date", "On approval", "Pass start date"],
            ["validTill", "Date", "Calculated", "Pass expiry date"],
            ["qrToken", "String", "Generated", "JWT QR token"],
            ["qrImage", "String", "Generated", "QR code data URL"],
            ["code16", "String", "Unique", "16-char hex verification code"],
            ["memberName", "String", "Family", "Family member name"],
            ["memberEmail", "String", "Family", "Family member email"],
            ["memberRelation", "String", "Enum", "Relation to user"],
            ["status", "String", "Enum", "pending/active/rejected/expired"],
        ],
        col_widths=[32, 22, 40, 96],
    )

    # ── 6. Business Logic ──
    pdf.add_page()
    pdf.add_heading("6. BUSINESS LOGIC")

    pdf.add_heading("I. Fare Calculation", level=2)
    pdf.add_paragraph(
        "The system uses the Haversine formula to calculate geographic distance between two bus stops "
        "using latitude and longitude coordinates. A road factor of 1.3x is applied for actual road distance."
    )
    pdf.add_bullet("Base Rate: Rs. 3 per kilometer")
    pdf.add_bullet("Minimum Fare: Rs. 200")
    pdf.add_bullet("Monthly Fare = max(distance x 3, 200)")
    pdf.add_bullet("Quarterly = Monthly x 2.4 (90 days)")
    pdf.add_bullet("Half-Yearly = Monthly x 4.0 (180 days)")
    pdf.add_bullet("Yearly = Monthly x 7.0 (365 days)")

    pdf.add_heading("II. Bus Stops", level=2)
    pdf.add_paragraph("The system supports 20 bus stops across Gujarat:")
    pdf.add_table(
        ["#", "Stop Name", "#", "Stop Name"],
        [
            ["1", "Ahmedabad (Paldi)", "11", "Navsari"],
            ["2", "Ahmedabad (Maninagar)", "12", "Valsad"],
            ["3", "Ahmedabad (Naroda)", "13", "Vapi"],
            ["4", "Gandhinagar", "14", "Rajkot"],
            ["5", "Nadiad", "15", "Jamnagar"],
            ["6", "Anand", "16", "Junagadh"],
            ["7", "Vadodara", "17", "Bhavnagar"],
            ["8", "Bharuch", "18", "Morbi"],
            ["9", "Surat", "19", "Mehsana"],
            ["10", "Palanpur", "20", "Himmatnagar"],
        ],
        col_widths=[10, 85, 10, 85],
    )

    pdf.add_heading("III. Pass Status Lifecycle", level=2)
    pdf.set_font("Courier", "", 10)
    pdf.set_text_color(0, 0, 0)
    lifecycle = [
        "  User Buys Pass         Admin Reviews          ",
        "  +----------+    +----------+----------+       ",
        "  |          |    |          |          |       ",
        "  | PENDING  +--->| APPROVE  |  REJECT  |       ",
        "  |          |    | (QR Gen) |          |       ",
        "  +----------+    +-----+----+----+-----+       ",
        "                        |         |             ",
        "                        v         v             ",
        "                    ACTIVE    REJECTED          ",
        "                      |                         ",
        "                      v                         ",
        "                   EXPIRED (past validTill)     ",
    ]
    for line in lifecycle:
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    pdf.add_heading("IV. QR Code Verification", level=2)
    pdf.add_paragraph(
        "When a pass is approved, a JWT token is generated containing pass ID, user ID, and expiry date. "
        "This token is encoded into a QR code image. A unique 16-character hexadecimal code is also "
        "generated for manual verification. Conductors can verify using either the QR scanner or the code."
    )

    # ── 7. Screenshots - User Side ──
    pdf.add_page()
    pdf.add_heading("7. SCREENSHOTS - USER SIDE")
    pdf.add_paragraph("User Portal URL: http://localhost:5173")
    pdf.add_paragraph("Production URL: https://bus-pass-management-system-user-frontend.onrender.com")

    user_ss = [(f, t, d) for f, t, d, s in SCREENSHOTS if s == "user"]
    for filename, title, desc in user_ss:
        if pdf.get_y() > 180:
            pdf.add_page()
        filepath = os.path.join(SS_DIR, filename)
        pdf.add_screenshot(filepath, title, desc)

    # ── 8. Screenshots - Admin Side ──
    pdf.add_page()
    pdf.add_heading("8. SCREENSHOTS - ADMIN SIDE")
    pdf.add_paragraph("Admin Portal URL: http://localhost:5174")
    pdf.add_paragraph("Production URL: https://bus-pass-management-system-1.onrender.com")

    admin_ss = [(f, t, d) for f, t, d, s in SCREENSHOTS if s == "admin"]
    for filename, title, desc in admin_ss:
        if pdf.get_y() > 180:
            pdf.add_page()
        filepath = os.path.join(SS_DIR, filename)
        pdf.add_screenshot(filepath, title, desc)

    # ── 9. Conclusion ──
    pdf.add_page()
    pdf.add_heading("9. CONCLUSION")
    pdf.add_paragraph(
        "The Bus Pass Management System successfully demonstrates a full-stack web application built "
        "using the MERN stack (MongoDB, Express.js, React.js, Node.js). The system provides a complete "
        "digital solution for bus pass management with the following key achievements:"
    )
    pdf.add_bullet("Digitized the entire bus pass lifecycle from purchase to verification.")
    pdf.add_bullet("Implemented secure JWT authentication with bcrypt password hashing and role-based access control.")
    pdf.add_bullet("Built intelligent fare calculation using the Haversine formula across 20 Gujarat bus stops.")
    pdf.add_bullet("Developed QR code-based verification for real-time pass validation by conductors.")
    pdf.add_bullet("Created family pass feature with bulk purchasing and auto account creation.")
    pdf.add_bullet("Designed responsive UI with Tailwind CSS for desktop and mobile devices.")
    pdf.add_bullet("Deployed on Render.com with separate user portal, admin portal, and backend API.")
    pdf.add_paragraph(
        "The project demonstrates proficiency in modern web development including RESTful API design, "
        "NoSQL database modeling, token-based authentication, QR code generation/scanning, and "
        "responsive UI design."
    )

    pdf.output(OUTPUT_PDF)
    print(f"PDF generated: {OUTPUT_PDF}")


# ═══════════════════════════════════════════════════════
#  DOCX Generation
# ═══════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_docx_table(doc, headers, rows):
    """Add a formatted table to the docx."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "000000")

    # Data rows
    for ri, row in enumerate(rows):
        row_cells = table.add_row().cells
        for ci, val in enumerate(row):
            cell = row_cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)
            if ri % 2 == 1:
                set_cell_shading(cell, "EBEBEB")

    doc.add_paragraph()


def add_docx_heading(doc, text, level=1):
    """Add heading with black color."""
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_docx_para(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add justified paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = alignment
    return p


def add_docx_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_docx_screenshot(doc, filepath, title, description):
    """Add screenshot with title and description."""
    p = doc.add_paragraph()
    run = p.add_run(f"Screenshot: {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)

    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"[Screenshot placeholder: {title}]")
        run2.font.size = Pt(11)
        run2.font.name = "Times New Roman"
        run2.italic = True
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption
    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"Figure: {title} - {description}")
    run3.font.size = Pt(11)
    run3.font.name = "Times New Roman"
    run3.font.color.rgb = RGBColor(60, 60, 60)
    run3.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()


def build_docx():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Cover Page ──
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GLS UNIVERSITY")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty of Computer Applications & Information Technology")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrated Master of Computer Applications [iMSc.IT]")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Programme")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Semester VIII Project Report for")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("221601801 FULL STACK DEVELOPMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"Bus Pass Management System"')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by:")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Group No.:\nHardik Sanjaybhai Parmar \u2013 202201619010054\nDhairy Pradipbhai Makadia \u2013 202201619010152")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    # ── Certificate Page ──
    doc.add_page_break()

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GLS UNIVERSITY")
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty of Computer Applications & IT iMSc.IT Programme")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ahmedabad")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CERTIFICATE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This is to certify that")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    for name in ["1. Hardik Sanjaybhai Parmar", "2. Dhairy Pradipbhai Makadia"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

    doc.add_paragraph()

    add_docx_para(
        doc,
        "Students of Semester-VIII iMSc.IT, FCAIT, GLS University have successfully "
        "completed the Project of 221601801 FULL STACK DEVELOPMENT as a part of their "
        "curriculum for the academic year 2025-2026."
    )

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Date of Submission:")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run2 = p.add_run("Project Guide - Name & Sign")
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"

    # ── Table of Contents ──
    doc.add_page_break()
    add_docx_heading(doc, "TABLE OF CONTENTS")

    toc = [
        ("1.", "Introduction", True),
        ("   I.", "Project Objective", False),
        ("   II.", "Purpose", False),
        ("   III.", "Modules", False),
        ("2.", "System Architecture", True),
        ("   I.", "Technology Stack", False),
        ("   II.", "Architecture Diagram", False),
        ("   III.", "Project Structure", False),
        ("3.", "Features & Functionality", True),
        ("4.", "API Endpoints", True),
        ("5.", "Database Schema Design", True),
        ("6.", "Business Logic", True),
        ("7.", "Screenshots - User Side", True),
        ("8.", "Screenshots - Admin Side", True),
        ("9.", "Conclusion", True),
    ]
    for num, title, is_main in toc:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.bold = is_main
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ── 1. Introduction ──
    doc.add_page_break()
    add_docx_heading(doc, "1. INTRODUCTION")

    add_docx_heading(doc, "I. Project Objective", level=2)
    add_docx_para(
        doc,
        "The Bus Pass Management System is a comprehensive full-stack web application designed to "
        "digitize and streamline the process of issuing, managing, and verifying bus passes. The system "
        "replaces traditional paper-based bus pass systems with a modern, efficient, and secure digital "
        "solution that benefits both commuters and transport authorities."
    )
    add_docx_para(
        doc,
        "The primary objective is to create a seamless experience for users to purchase bus passes online, "
        "enable administrators to manage pass approvals efficiently, and provide conductors with real-time "
        "QR-based verification capabilities. The system supports individual and family pass purchases with "
        "dynamic fare calculation based on distance between bus stops."
    )

    add_docx_heading(doc, "II. Purpose", level=2)
    add_docx_para(doc, "The purpose of this system is to modernize public transportation services by providing:")
    add_docx_bullet(doc, "Digital bus passes with QR code verification")
    add_docx_bullet(doc, "Online application and approval system")
    add_docx_bullet(doc, "Family pass management with auto-account creation")
    add_docx_bullet(doc, "Real-time pass tracking and status updates")
    add_docx_bullet(doc, "Role-based access for Admin, Superadmin, and Conductor")
    add_docx_para(
        doc,
        "This system reduces waiting time at bus pass counters and ensures a smooth and secure "
        "experience for users and administrators."
    )

    add_docx_heading(doc, "III. Modules", level=2)

    add_docx_heading(doc, "1. User Authentication Module", level=3)
    add_docx_bullet(doc, "User registration and login with form validation")
    add_docx_bullet(doc, "Secure authentication using JWT (JSON Web Tokens)")
    add_docx_bullet(doc, "Password encryption using Bcrypt")
    add_docx_bullet(doc, "Forgot and reset password functionality via email")

    add_docx_heading(doc, "2. Individual Pass Module", level=3)
    add_docx_bullet(doc, "Select source and destination from 20 bus stops")
    add_docx_bullet(doc, "Dynamic fare calculation using Haversine formula")
    add_docx_bullet(doc, "Choose pass type (Monthly, Quarterly, Half-Yearly, Yearly)")
    add_docx_bullet(doc, "Submit pass request for admin approval")

    add_docx_heading(doc, "3. Family Pass Module", level=3)
    add_docx_bullet(doc, "Add multiple family members (up to 10)")
    add_docx_bullet(doc, "Auto-create user accounts for family members")
    add_docx_bullet(doc, "Purchase passes for entire family in one transaction")
    add_docx_bullet(doc, "3-step wizard: Add Members, Route & Pass, Review")

    add_docx_heading(doc, "4. Dashboard Module", level=3)
    add_docx_bullet(doc, "View all passes with status (Active, Pending, Expired, Rejected)")
    add_docx_bullet(doc, "Display QR codes for active passes")
    add_docx_bullet(doc, "Pass details with verification code")
    add_docx_bullet(doc, "Responsive design for mobile and desktop")

    add_docx_heading(doc, "5. Admin Module", level=3)
    add_docx_bullet(doc, "Approve or reject pending pass applications")
    add_docx_bullet(doc, "View passes filtered by status")
    add_docx_bullet(doc, "Statistics dashboard (Total, Active, Pending counts)")
    add_docx_bullet(doc, "Superadmin can manage other admin accounts")

    add_docx_heading(doc, "6. Verification Module", level=3)
    add_docx_bullet(doc, "QR code scanning using device camera")
    add_docx_bullet(doc, "Manual verification using 16-character hex code")
    add_docx_bullet(doc, "Real-time validation with pass details display")

    # ── 2. System Architecture ──
    doc.add_page_break()
    add_docx_heading(doc, "2. SYSTEM ARCHITECTURE")

    add_docx_heading(doc, "I. Technology Stack", level=2)
    add_docx_table(doc,
        ["Component", "Technology"],
        [
            ["Frontend (User)", "React.js 19.x + Tailwind CSS + Vite"],
            ["Frontend (Admin)", "React.js 19.x + Tailwind CSS + Vite"],
            ["Backend", "Node.js + Express.js 5.x"],
            ["Database", "MongoDB Atlas (Mongoose 9.x ODM)"],
            ["Authentication", "JSON Web Token (JWT)"],
            ["Password Security", "Bcrypt.js"],
            ["QR Code", "qrcode library + html5-qrcode scanner"],
            ["Email Service", "Nodemailer (Gmail SMTP)"],
            ["Routing", "React Router 7.x"],
            ["Deployment", "Render.com (Cloud Platform)"],
        ],
    )

    add_docx_heading(doc, "II. Project Structure", level=2)
    structure = (
        "bus-pass-management-system/\n"
        "|-- backend/\n"
        "|   |-- controllers/     # Business logic handlers\n"
        "|   |-- models/          # Mongoose schemas\n"
        "|   |-- routes/          # API route definitions\n"
        "|   |-- middleware/      # Auth guards, HTTP logging\n"
        "|   |-- utils/           # Token, QR, email, logger\n"
        "|   |-- index.js         # Express server entry point\n"
        "|\n"
        "|-- frontend/            # Admin Portal (React + Vite)\n"
        "|   |-- src/auth/        # Login, Register, auth service\n"
        "|   |-- src/pages/       # Admin Dashboard\n"
        "|\n"
        "|-- user-frontend/       # User Portal (React + Vite)\n"
        "|   |-- src/auth/        # Login, Register, ForgotPassword\n"
        "|   |-- src/pages/       # Dashboard, FamilyPass, Profile\n"
        "|   |-- src/components/  # Navbar, Footer\n"
        "|   |-- src/layout/      # RootLayout wrapper"
    )
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.size = Pt(10)
    run.font.name = "Courier New"
    run.font.color.rgb = RGBColor(0, 0, 0)

    add_docx_heading(doc, "III. Database Models", level=2)
    add_docx_para(doc, "The system uses three MongoDB collections: Users, Admins, and BusPasses. Each collection is defined with Mongoose schemas providing validation, middleware hooks (password hashing on save), and relationship references.")

    # ── 3. Features ──
    doc.add_page_break()
    add_docx_heading(doc, "3. FEATURES & FUNCTIONALITY")

    add_docx_heading(doc, "I. User Features", level=2)
    add_docx_bullet(doc, "User Registration with name, email, phone (10-digit), DOB, and strong password validation")
    add_docx_bullet(doc, "JWT-based login with 1-day token expiry")
    add_docx_bullet(doc, "Email-based password reset with 6-character code (15-min expiry)")
    add_docx_bullet(doc, "Individual pass purchase with real-time fare display for 4 pass types")
    add_docx_bullet(doc, "Family pass purchase for up to 10 members with auto-account creation")
    add_docx_bullet(doc, "View pass history with QR codes, verification codes, and validity dates")
    add_docx_bullet(doc, "Responsive design for desktop and mobile devices")

    add_docx_heading(doc, "II. Admin Features", level=2)
    add_docx_bullet(doc, "Pass approval/rejection workflow with QR code generation on approval")
    add_docx_bullet(doc, "Camera-based QR scanner for real-time pass verification")
    add_docx_bullet(doc, "Manual verification using 16-character hex code")
    add_docx_bullet(doc, "Pass filtering by status (Pending, Active, Rejected, All)")
    add_docx_bullet(doc, "Role-based access: Admin (manage passes), Conductor (verify only), Superadmin (full control)")

    add_docx_heading(doc, "III. Authentication & Security", level=2)
    add_docx_bullet(doc, "JWT token-based stateless authentication with 1-day expiration")
    add_docx_bullet(doc, "Bcrypt password hashing with salt rounds")
    add_docx_bullet(doc, "Role-based middleware for users, admins, superadmins, and conductors")
    add_docx_bullet(doc, "CORS whitelisting for development and production origins")
    add_docx_bullet(doc, "Server-side input validation for all endpoints")

    # ── 4. API Endpoints ──
    doc.add_page_break()
    add_docx_heading(doc, "4. API ENDPOINTS")

    add_docx_heading(doc, "User Authentication API", level=2)
    add_docx_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/user/register", "Register new user"],
            ["POST", "/api/user/login", "User login, returns JWT"],
            ["POST", "/api/user/forgot-password", "Send reset code to email"],
            ["POST", "/api/user/reset-password", "Reset password with code"],
            ["GET", "/api/user/profile", "Get user profile"],
            ["PUT", "/api/user/change-password", "Change password"],
        ],
    )

    add_docx_heading(doc, "Bus Pass API", level=2)
    add_docx_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/user/bus-pass/stops", "Get 20 bus stops list"],
            ["GET", "/api/user/bus-pass/fare", "Calculate fare for route"],
            ["POST", "/api/user/bus-pass/buy", "Purchase individual pass"],
            ["POST", "/api/user/bus-pass/buy-family", "Purchase family passes"],
            ["GET", "/api/user/bus-pass/my-passes", "Get user's passes"],
            ["POST", "/api/user/bus-pass/verify", "Verify pass (QR/code)"],
        ],
    )

    add_docx_heading(doc, "Admin API", level=2)
    add_docx_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/admin/login", "Admin login"],
            ["POST", "/api/admin/register", "Register admin"],
            ["GET", "/api/admin/bus-passes", "Get passes by status"],
            ["PATCH", "/api/admin/bus-passes/:id/approve", "Approve pass"],
            ["PATCH", "/api/admin/bus-passes/:id/reject", "Reject pass"],
            ["POST", "/api/admin/bus-passes/verify", "Verify pass"],
        ],
    )

    # ── 5. Database Schema ──
    doc.add_page_break()
    add_docx_heading(doc, "5. DATABASE SCHEMA DESIGN")

    add_docx_heading(doc, "User Collection", level=2)
    add_docx_table(doc,
        ["Field", "Type", "Description"],
        [
            ["name", "String", "User's full name (required)"],
            ["email", "String", "Unique email address (required)"],
            ["phone", "String", "10-digit phone number (required)"],
            ["dateOfBirth", "Date", "Date of birth (required)"],
            ["password", "String", "Bcrypt hashed password (required)"],
            ["isActive", "Boolean", "Account active status (default: true)"],
        ],
    )

    add_docx_heading(doc, "Admin Collection", level=2)
    add_docx_table(doc,
        ["Field", "Type", "Description"],
        [
            ["name", "String", "Admin's full name (required)"],
            ["email", "String", "Unique email address (required)"],
            ["password", "String", "Bcrypt hashed password (required)"],
            ["role", "String", "admin / superadmin / conductor"],
            ["isActive", "Boolean", "Account active status (default: true)"],
        ],
    )

    add_docx_heading(doc, "BusPass Collection", level=2)
    add_docx_table(doc,
        ["Field", "Type", "Description"],
        [
            ["user", "ObjectId", "Reference to User (pass owner)"],
            ["passType", "String", "monthly / quarterly / half-yearly / yearly"],
            ["price", "Number", "Calculated fare amount"],
            ["source", "String", "Source bus stop name"],
            ["destination", "String", "Destination bus stop name"],
            ["validFrom", "Date", "Pass start date (set on approval)"],
            ["validTill", "Date", "Pass expiry date (calculated)"],
            ["qrToken", "String", "JWT encoded QR token"],
            ["qrImage", "String", "QR code as data URL"],
            ["code16", "String", "16-char hex verification code"],
            ["memberName", "String", "Family member name"],
            ["memberRelation", "String", "Relation to primary user"],
            ["status", "String", "pending / active / rejected / expired"],
        ],
    )

    # ── 6. Business Logic ──
    doc.add_page_break()
    add_docx_heading(doc, "6. BUSINESS LOGIC")

    add_docx_heading(doc, "I. Fare Calculation", level=2)
    add_docx_para(
        doc,
        "The system uses the Haversine formula to calculate geographic distance between two bus stops "
        "using latitude and longitude coordinates. A road factor of 1.3x is applied for actual road distance."
    )
    add_docx_bullet(doc, "Base Rate: Rs. 3 per kilometer")
    add_docx_bullet(doc, "Minimum Fare: Rs. 200")
    add_docx_bullet(doc, "Monthly Fare = max(distance x 3, 200)")
    add_docx_bullet(doc, "Quarterly = Monthly x 2.4 (90 days)")
    add_docx_bullet(doc, "Half-Yearly = Monthly x 4.0 (180 days)")
    add_docx_bullet(doc, "Yearly = Monthly x 7.0 (365 days)")

    add_docx_heading(doc, "II. Bus Stops", level=2)
    add_docx_para(doc, "The system supports 20 bus stops across Gujarat state:")
    add_docx_table(doc,
        ["#", "Stop Name", "#", "Stop Name"],
        [
            ["1", "Ahmedabad (Paldi)", "11", "Navsari"],
            ["2", "Ahmedabad (Maninagar)", "12", "Valsad"],
            ["3", "Ahmedabad (Naroda)", "13", "Vapi"],
            ["4", "Gandhinagar", "14", "Rajkot"],
            ["5", "Nadiad", "15", "Jamnagar"],
            ["6", "Anand", "16", "Junagadh"],
            ["7", "Vadodara", "17", "Bhavnagar"],
            ["8", "Bharuch", "18", "Morbi"],
            ["9", "Surat", "19", "Mehsana"],
            ["10", "Palanpur", "20", "Himmatnagar"],
        ],
    )

    add_docx_heading(doc, "III. Pass Status Lifecycle", level=2)
    add_docx_para(doc, "PENDING --> APPROVED (QR Generated) --> ACTIVE --> EXPIRED")
    add_docx_para(doc, "PENDING --> REJECTED")

    add_docx_heading(doc, "IV. QR Code Verification", level=2)
    add_docx_para(
        doc,
        "When a pass is approved, a JWT token is generated containing pass ID, user ID, and expiry date. "
        "This token is encoded into a QR code image. A unique 16-character hexadecimal code is also "
        "generated for manual verification. Conductors can verify using either the QR scanner or the code."
    )

    # ── 7. Screenshots - User Side ──
    doc.add_page_break()
    add_docx_heading(doc, "7. SCREENSHOTS - USER SIDE")
    add_docx_para(doc, "User Portal URL: http://localhost:5173")
    add_docx_para(doc, "Production: https://bus-pass-management-system-user-frontend.onrender.com")

    user_ss = [(f, t, d) for f, t, d, s in SCREENSHOTS if s == "user"]
    for filename, title, desc in user_ss:
        filepath = os.path.join(SS_DIR, filename)
        add_docx_screenshot(doc, filepath, title, desc)

    # ── 8. Screenshots - Admin Side ──
    doc.add_page_break()
    add_docx_heading(doc, "8. SCREENSHOTS - ADMIN SIDE")
    add_docx_para(doc, "Admin Portal URL: http://localhost:5174")
    add_docx_para(doc, "Production: https://bus-pass-management-system-1.onrender.com")

    admin_ss = [(f, t, d) for f, t, d, s in SCREENSHOTS if s == "admin"]
    for filename, title, desc in admin_ss:
        filepath = os.path.join(SS_DIR, filename)
        add_docx_screenshot(doc, filepath, title, desc)

    # ── 9. Conclusion ──
    doc.add_page_break()
    add_docx_heading(doc, "9. CONCLUSION")
    add_docx_para(
        doc,
        "The Bus Pass Management System successfully demonstrates a full-stack web application built "
        "using the MERN stack (MongoDB, Express.js, React.js, Node.js). The system provides a complete "
        "digital solution for bus pass management with the following key achievements:"
    )
    add_docx_bullet(doc, "Digitized the entire bus pass lifecycle from purchase to verification.")
    add_docx_bullet(doc, "Implemented secure JWT authentication with bcrypt password hashing and role-based access control.")
    add_docx_bullet(doc, "Built intelligent fare calculation using the Haversine formula across 20 Gujarat bus stops.")
    add_docx_bullet(doc, "Developed QR code-based verification for real-time pass validation by conductors.")
    add_docx_bullet(doc, "Created family pass feature with bulk purchasing and auto account creation.")
    add_docx_bullet(doc, "Designed responsive UI with Tailwind CSS for desktop and mobile devices.")
    add_docx_bullet(doc, "Deployed on Render.com with separate user portal, admin portal, and backend API.")
    add_docx_para(
        doc,
        "The project demonstrates proficiency in modern web development including RESTful API design, "
        "NoSQL database modeling, token-based authentication, QR code generation/scanning, and "
        "responsive UI design."
    )

    doc.save(OUTPUT_DOCX)
    print(f"DOCX generated: {OUTPUT_DOCX}")


# ═══════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating documentation...")
    print(f"Screenshots directory: {SS_DIR}")
    print(f"Screenshots found: {len([f for f in os.listdir(SS_DIR) if f.endswith('.png')]) if os.path.exists(SS_DIR) else 0}")
    print()
    build_pdf()
    build_docx()
    print("\nDone!")
